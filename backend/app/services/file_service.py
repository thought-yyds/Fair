from pathlib import Path
from datetime import datetime
from uuid import uuid4  # 新增：用于生成唯一文件名（解决重名和路径攻击）
import re  # 新增：用于完善句子分割（匹配多种句末标点）
from fastapi import UploadFile, HTTPException
from docx import Document
from pdf2docx import Converter

from app.config import settings
from app.models import db_models, schemas
from app.utils.helpers import clean_text


def allowed_file(filename: str) -> bool:
    """检查文件是否为允许的类型（docx/pdf）"""
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in settings.ALLOWED_EXTENSIONS


def save_uploaded_file(file: UploadFile) -> tuple[db_models.Article, Path, Path]:
    """
    保存上传的文件，并生成批注文件路径
    :param file: 上传的文件（FastAPI的UploadFile对象）
    :return: (Article数据库模型, 原始文件路径, 批注文件路径)
    """
    # 1. 验证文件类型
    if not allowed_file(file.filename):
        raise HTTPException(status_code=400, detail=f"不支持的文件类型：{file.filename.split('.')[-1]}，仅支持docx/pdf")

    # 2. 定义文件路径（核心修改：用UUID生成唯一文件名，避免重名和路径攻击）
    original_filename = file.filename  # 保留原始文件名（给用户看）
    file_ext = original_filename.rsplit('.', 1)[1].lower()  # 提取文件扩展名（docx/pdf）
    unique_filename = f"{uuid4()}.{file_ext}"  # 生成唯一文件名（如：a1b2c3-d4e5f6.docx）
    original_path = settings.UPLOADS_DIR / unique_filename  # 原始文件存储路径（唯一文件名）

    # 3. 检查是否已存在同名文件（用Path方法替代os，风格统一）
    if original_path.exists():
        raise HTTPException(status_code=400, detail=f"文件已上传，请重新上传")

    # 4. 保存原始文件（二进制写入）
    with open(original_path, "wb") as f:
        f.write(file.file.read())

    # 5. 生成批注文件路径 + 处理文件转换（核心修改：完善PDF转docx逻辑）
    annotated_filename = f"annotated_{uuid4()}.docx"  # 批注文件也用唯一文件名
    annotated_path = settings.UPLOADS_DIR / annotated_filename

    if file_ext == 'pdf':
        # 5.1 PDF文件：用pdf2docx转换为docx（解决原空文件问题）
        try:
            cv = Converter(str(original_path))  # 初始化PDF转换器
            cv.convert(str(annotated_path), start=0, end=None)  # 转换全部页面
            cv.close()
        except Exception as e:
            # 转换失败：清理临时文件（避免残留）+ 抛友好错误
            if original_path.exists():
                original_path.unlink()  # Path方法：删除文件
            if annotated_path.exists():
                annotated_path.unlink()
            raise HTTPException(status_code=500, detail=f"PDF转Word失败：{str(e)}，请尝试直接上传docx文件")
    else:
        # 5.2 DOCX文件：直接复制原始文件作为批注文件（保持原有逻辑）
        with open(original_path, "rb") as f1, open(annotated_path, "wb") as f2:
            f2.write(f1.read())

    # 6. 创建Article数据库记录（注意：name存原始文件名，path存唯一路径）
    article = db_models.Article(
        name=original_filename,  # 前端展示用：用户上传的原始文件名
        original_path=str(original_path),  # 后端存储用：唯一文件名路径
        annotated_path=str(annotated_path),
        upload_time=datetime.utcnow(),
        status="待审查",
        review_progress=0
    )

    return article, original_path, annotated_path


def extract_sentences_from_docx(docx_path: Path) -> list[str]:
    """
    从docx文件中提取句子（核心修改：完善句子分割逻辑，支持。！？）
    :param docx_path: docx文件路径（Path对象）
    :return: 清洗后的句子列表
    """
    # 验证路径有效性（用Path方法替代os，更严谨）
    if not docx_path.exists() or docx_path.suffix.lower() != '.docx':
        raise HTTPException(status_code=400, detail=f"无效的docx文件：{str(docx_path)}")

    doc = Document(docx_path)
    sentences = []

    for paragraph in doc.paragraphs:
        clean_paragraph = clean_text(paragraph.text)
        if not clean_paragraph:
            continue  # 跳过空段落

        # 核心修改：用正则匹配“。”“！”“？”作为句末标点，分割更准确
        para_sentences = re.split(r'[。！？]', clean_paragraph)
        # 二次清洗：过滤空字符串 + 去除前后空格
        para_sentences = [clean_text(s) for s in para_sentences if clean_text(s)]
        sentences.extend(para_sentences)

    return sentences


def read_full_doc_content(docx_path: Path) -> str:
    """
    读取docx文件的完整文本内容（保留段落结构，用\n分隔）
    注：PDF文件已在save_uploaded_file中转为docx，故只需处理docx
    """
    if not docx_path.exists() or docx_path.suffix.lower() != ".docx":
        raise HTTPException(
            status_code=400,
            detail=f"无效的文档文件：{str(docx_path)}，仅支持docx格式"
        )

    doc = Document(docx_path)
    # 拼接所有非空段落，用\n保留段落换行（前端可直接按\n渲染换行）
    full_content = "\n".join([
        para.text.strip()
        for para in doc.paragraphs
        if para.text.strip()  # 过滤空段落
    ])
    return full_content


def extract_sentences_with_position(full_content: str) -> list[dict]:
    """
    从完整文本中提取句子，并记录每个句子的「起始索引」和「结束索引」（用于前端高亮）
    返回格式：[{"content": "句子内容", "start_idx": 0, "end_idx": 15}, ...]
    """
    sentences = []
    current_pos = 0  # 当前文本的起始位置指针
    text_length = len(full_content)

    # 用正则匹配句末标点（支持。！？；，覆盖常见中文标点）
    # 正则说明：匹配非标点字符+句末标点，避免拆分不完整
    sentence_pattern = re.compile(r'[^。！？；，]*[。！？；，]')
    matches = sentence_pattern.finditer(full_content)

    for match in matches:
        sentence_text = match.group().strip()
        if not sentence_text:  # 过滤空句子
            current_pos = match.end()
            continue

        # 记录句子在完整文本中的起始/结束索引
        start_idx = match.start()
        end_idx = match.end()

        sentences.append({
            "content": sentence_text,
            "start_idx": start_idx,
            "end_idx": end_idx
        })

        current_pos = end_idx  # 更新指针到当前句子末尾

    # 处理最后一个没有标点的句子（如文本末尾的短句）
    if current_pos < text_length:
        remaining_text = full_content[current_pos:].strip()
        if remaining_text:
            sentences.append({
                "content": remaining_text,
                "start_idx": current_pos,
                "end_idx": text_length  # 结束索引=文本总长度
            })

    return sentences
